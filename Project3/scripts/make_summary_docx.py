# -*- coding: utf-8 -*-
"""Sinh file Word tóm tắt đồ án (ngắn gọn cho giảng viên)."""
import os
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

ROOT = os.path.dirname(os.path.dirname(os.path.abspath(__file__)))
CHART = os.path.join(ROOT, "analysis", "charts", "model_comparison.png")
OUT = os.path.join(ROOT, "reports", "TomTat_DoAn.docx")

doc = Document()
# Font mặc định
style = doc.styles["Normal"]
style.font.name = "Times New Roman"
style.font.size = Pt(12)

def h(text, level=1):
    p = doc.add_heading(text, level=level)
    return p

def para(text, bold=False, italic=False, align=None, size=None):
    p = doc.add_paragraph()
    r = p.add_run(text)
    r.bold = bold; r.italic = italic
    if size: r.font.size = Pt(size)
    if align is not None: p.alignment = align
    return p

def bullet(text):
    p = doc.add_paragraph(style="List Bullet")
    p.add_run(text)
    return p

# ===== Tiêu đề =====
t = doc.add_paragraph()
t.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = t.add_run("TÓM TẮT ĐỒ ÁN TỐT NGHIỆP")
r.bold = True; r.font.size = Pt(16)

s = doc.add_paragraph()
s.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = s.add_run("Hệ thống phát hiện xâm nhập cho web server dựa trên phân tích log,\n"
              "kết hợp phương pháp dựa trên dấu hiệu (signature) và phát hiện bất thường (AI)")
r.italic = True; r.font.size = Pt(12)

info = doc.add_paragraph()
info.alignment = WD_ALIGN_PARAGRAPH.CENTER
info.add_run("Sinh viên: Trần Công Thành   |   GVHD: TS. Nguyễn Hữu Đức   |   HUST, 06/2026").font.size = Pt(11)

# ===== 1. Đồ án này là gì =====
h("1. Đồ án này là gì?", 1)
para("Đồ án xây dựng một hệ thống phát hiện tấn công (IDS) nhằm vào máy chủ web. "
     "Hệ thống đọc trực tiếp nhật ký truy cập (access log) của web server, và với mỗi "
     "yêu cầu HTTP/HTTPS, nó kết luận yêu cầu đó là BÌNH THƯỜNG hay TẤN CÔNG, đồng thời "
     "xếp hạng mức độ nghiêm trọng để hỗ trợ người quản trị xử lý. Điểm cốt lõi là hệ "
     "thống không dùng một phương pháp đơn lẻ mà KẾT HỢP ba tầng phát hiện bổ trợ nhau, "
     "qua đó vừa chính xác với tấn công đã biết, vừa có khả năng phát hiện tấn công mới.")

# ===== 2. Bài toán & động lực =====
h("2. Bài toán và động lực", 1)
para("Các nghiên cứu trước thường chỉ theo một hướng: hoặc dùng luật/dấu hiệu (chính xác "
     "nhưng bỏ lọt tấn công mới), hoặc chỉ dùng học máy có giám sát (cần nhãn, khó tổng "
     "quát hóa), hoặc chỉ dùng phát hiện bất thường (bắt được cái mới nhưng nhiều báo "
     "động giả). Đồ án giải quyết khoảng trống này bằng một kiến trúc lai ghép ba tầng "
     "có cơ chế đồng thuận thông minh để tận dụng điểm mạnh của cả ba hướng.")

# ===== 3. Giải pháp đề xuất =====
h("3. Giải pháp đề xuất: kiến trúc 3 tầng", 1)
para("Mỗi yêu cầu sau khi được phân tích và giải mã URL nhiều lớp sẽ được biểu diễn thành "
     "một véc-tơ 22 đặc trưng, rồi đưa qua ba tầng:")
bullet("Tầng 1 – Dấu hiệu (Regex/luật): so khớp các mẫu tấn công đã biết như SQL Injection, "
       "XSS, duyệt thư mục, chèn lệnh… Rất chính xác nhưng độ phủ thấp.")
bullet("Tầng 2 – Học có giám sát (Random Forest): học từ dữ liệu đã gán nhãn, cho ra xác "
       "suất tấn công. Đây là mô hình mạnh nhất khi đứng riêng.")
bullet("Tầng 3 – Phát hiện bất thường không giám sát (Local Outlier Factor): chỉ học từ "
       "lưu lượng SẠCH, coi mọi sai lệch lớn là bất thường. Nhờ vậy bắt được tấn công mới "
       "mà không cần nhãn tấn công.")
para("Cơ chế đồng thuận thông minh (Smart Consensus): hệ thống báo động khi có luật khớp, "
     "HOẶC Random Forest đủ tự tin, HOẶC tầng bất thường phát hiện VÀ được Random Forest "
     "xác nhận ở mức tối thiểu. Việc để Tầng 2 xác thực Tầng 3 giúp giảm mạnh báo động giả "
     "mà vẫn giữ độ phủ cao.")

# ===== 4. Dữ liệu =====
h("4. Dữ liệu sử dụng", 1)
para("Đồ án xây dựng tập dữ liệu cân bằng gồm 111.065 dòng log, một nửa tấn công và một "
     "nửa bình thường, tổng hợp từ: tập chuẩn CSIC 2010, các mẫu tấn công thực tế từ kho "
     "PayloadsAllTheThings, và lưu lượng sinh tự động. Tầng học có giám sát được huấn luyện "
     "trên 66.984 mẫu có nhãn; tầng bất thường chỉ huấn luyện trên 34.941 mẫu SẠCH; phần "
     "đánh giá gồm 28.767 mẫu (47,43% là tấn công).")

# ===== 5. Kết quả =====
h("5. Kết quả chính", 1)
para("Bảng dưới đây so sánh từng tầng và các phương án kết hợp trên tập đánh giá (đơn vị: %). "
     "Chỉ số quan trọng nhất với IDS là F2 (ưu tiên độ phủ, vì bỏ lọt tấn công nguy hiểm hơn "
     "báo động giả).")

rows = [
    ("Cấu hình", "Precision", "Recall", "F1", "F2"),
    ("Tầng 1 – Regex", "99.61", "24.07", "38.77", "28.37"),
    ("Tầng 2 – Random Forest", "94.80", "94.76", "94.78", "94.77"),
    ("Tầng 3 – LOF", "88.28", "95.13", "91.57", "93.67"),
    ("Simple Voting", "87.78", "97.96", "92.59", "95.74"),
    ("Smart Consensus (đề xuất)", "91.19", "97.62", "94.30", "96.26"),
]
table = doc.add_table(rows=len(rows), cols=5)
table.style = "Light Grid Accent 1"
for i, row in enumerate(rows):
    for j, val in enumerate(row):
        cell = table.cell(i, j)
        cell.text = val
        for p in cell.paragraphs:
            for run in p.runs:
                run.font.size = Pt(11)
                if i == 0 or (i == len(rows) - 1):
                    run.bold = True

para("")
para("Nhận xét: Random Forest là mô hình đơn tốt nhất (F1 = 94,78%). Phương án Smart "
     "Consensus đạt F2 cao nhất toàn hệ thống (96,26%) với độ phủ 97,62% — tức bắt được "
     "khoảng 98/100 tấn công — trong khi vẫn giữ độ chính xác hợp lý (91,19%). So với "
     "Simple Voting, cơ chế đồng thuận giúp giảm đáng kể báo động giả.", italic=True)

# Biểu đồ
if os.path.exists(CHART):
    doc.add_picture(CHART, width=Inches(6.0))
    cap = doc.paragraphs[-1]
    doc.add_paragraph().alignment = WD_ALIGN_PARAGRAPH.CENTER
    c = doc.add_paragraph()
    c.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = c.add_run("Hình 1. So sánh các tầng và phương án kết hợp trên 4 chỉ số.")
    r.italic = True; r.font.size = Pt(10)

# ===== 6. Đóng góp nổi bật =====
h("6. Điểm nổi bật", 1)
bullet("Kết hợp đồng thời 3 hướng phát hiện trong một hệ thống, với cơ chế đồng thuận "
       "để Tầng 2 xác thực Tầng 3.")
bullet("Hoạt động trực tiếp trên log Apache thật, có giải mã nhiều lớp chống né tránh và "
       "hỗ trợ phân tích cả thân yêu cầu POST.")
bullet("Tuân thủ đúng nguyên lý học một lớp (chỉ huấn luyện trên dữ liệu sạch) — yếu tố "
       "giúp F1 của tầng bất thường tăng từ ~33% lên ~91%.")
bullet("Cảnh báo có thể giải thích (biết tầng nào kích hoạt) và có dashboard giám sát "
       "thời gian thực.")

# ===== 7. Hạn chế & hướng phát triển =====
h("7. Hạn chế và hướng phát triển", 1)
para("Hạn chế: mô hình huấn luyện chủ yếu trên CSIC 2010 nên có thể báo động giả nhiều "
     "hơn khi triển khai trên ứng dụng có hành vi rất khác; hệ thống mới dừng ở phát hiện "
     "và cảnh báo, chưa chặn. Hướng phát triển: đánh giá trên dữ liệu thực tế đa dạng hơn, "
     "học bán giám sát/trực tuyến để thích nghi theo thời gian, và tích hợp với tường lửa "
     "ứng dụng web (WAF) để chuyển từ phát hiện sang ngăn chặn.")

# ===== 8. Công nghệ =====
h("8. Công nghệ sử dụng", 1)
para("Python, scikit-learn (Random Forest, Logistic Regression, Isolation Forest, "
     "One-Class SVM, Local Outlier Factor), biểu thức chính quy cho tầng dấu hiệu, và "
     "Streamlit cho dashboard giám sát thời gian thực.")

doc.save(OUT)
print("Saved:", OUT)
